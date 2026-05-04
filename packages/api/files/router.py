"""File summarization — accepts an uploaded file, extracts text via
the right library for its format, and returns an LLM-grounded summary.

Replaces `smartMockSummary` in the frontend chat for non-textual
files (PDF, DOCX). Plain text + markdown are also supported here so
the frontend can route everything through one path.

Image OCR is out of scope for this iteration: images return a clear
"needs OCR" state rather than a fake summary.

Limits:
- 10 MB hard cap per upload (matches the chat-page UI limit)
- Extracted text trimmed to 64 KB before the LLM call so the prompt
  doesn't blow past the context window or rate-limit our budget
"""

from __future__ import annotations

import io
import logging

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from pydantic import BaseModel

from auth.deps import get_current_agent_optional
from agents.llm_providers import call_llm, get_active_provider
from db import AgentRow

router = APIRouter()
log = logging.getLogger("files.summarize")


MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_TEXT_BYTES = 64 * 1024            # 64 KB into the prompt


class SummarizeResponse(BaseModel):
    name: str
    mime_type: str
    size_bytes: int
    extracted_chars: int
    summary: str
    extractor: str  # "pypdf" | "python-docx" | "text" | "fallback"
    truncated: bool


def _looks_like_text(mime: str, name: str) -> bool:
    m = (mime or "").lower()
    if m.startswith("text/"):
        return True
    if m in {
        "application/json",
        "application/xml",
        "application/javascript",
        "application/x-yaml",
        "application/x-typescript",
        "application/sql",
    }:
        return True
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    return ext in {
        "md", "txt", "csv", "tsv", "json", "yaml", "yml", "xml", "html",
        "htm", "css", "scss", "js", "jsx", "ts", "tsx", "py", "rb", "go",
        "rs", "java", "kt", "swift", "c", "h", "cpp", "hpp", "sh", "bash",
        "zsh", "sql", "ini", "cfg", "toml", "log", "rst", "tex",
    }


def _looks_like_pdf(mime: str, name: str) -> bool:
    return (mime or "").lower() == "application/pdf" or name.lower().endswith(".pdf")


def _looks_like_docx(mime: str, name: str) -> bool:
    m = (mime or "").lower()
    if m == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return True
    return name.lower().endswith(".docx")


def _looks_like_image(mime: str, name: str) -> bool:
    if (mime or "").lower().startswith("image/"):
        return True
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    return ext in {"png", "jpg", "jpeg", "gif", "webp", "bmp", "tiff", "tif"}


def _extract_pdf(blob: bytes) -> str:
    """Pull text out of a PDF. Page-by-page, joined with double newline.
    Catches per-page errors so one bad page doesn't kill the whole pull."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(blob))
    chunks: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            chunks.append(text.strip())
    return "\n\n".join(chunks)


def _extract_docx(blob: bytes) -> str:
    """Pull body text out of a .docx. Paragraphs joined with newlines.
    Doesn't include headers/footers/comments — body is enough for summary."""
    from docx import Document

    doc = Document(io.BytesIO(blob))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tables: flatten cell text into rows so structured data shows up too.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _build_summary_prompt(name: str, mime: str, content: str) -> str:
    return (
        "You are summarizing a file the user just uploaded. Read the "
        "content carefully and produce a SPECIFIC summary that names "
        "actual people, dates, decisions, numbers, and action items "
        "found in the file. Do not give a generic 'this document "
        "discusses…' summary.\n\n"
        f"File name: {name}\n"
        f"MIME type: {mime or 'unknown'}\n\n"
        "----- BEGIN FILE CONTENT -----\n"
        f"{content}\n"
        "----- END FILE CONTENT -----\n\n"
        "Reply in markdown with:\n"
        "- A one-line gist (be specific to this file)\n"
        "- 3-6 bullets covering the concrete things actually in the file "
        "(people, dates, decisions, action items, numbers)\n"
        "- An 'Open questions' line only if the file has unresolved items\n\n"
        "If the content is truncated or unintelligible, say so plainly "
        "instead of inventing details."
    )


@router.post("/summarize", response_model=SummarizeResponse)
async def summarize_file(
    file: UploadFile = File(...),
    caller: AgentRow | None = Depends(get_current_agent_optional),
) -> SummarizeResponse:
    """Extract text + summarize. PDF, DOCX, and any text-like format are
    supported. Images return an 'unsupported — needs OCR' summary."""
    # Read once into memory; we cap before extraction to keep the
    # uvicorn worker honest. Streaming would be nicer for huge PDFs
    # but we explicitly cap upload size at the UI layer too.
    blob = await file.read()
    if len(blob) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="file_too_large")
    if not blob:
        raise HTTPException(status_code=400, detail="empty_file")

    name = file.filename or "uploaded_file"
    mime = file.content_type or ""

    extractor: str
    text: str
    try:
        if _looks_like_pdf(mime, name):
            extractor = "pypdf"
            text = _extract_pdf(blob)
        elif _looks_like_docx(mime, name):
            extractor = "python-docx"
            text = _extract_docx(blob)
        elif _looks_like_text(mime, name):
            extractor = "text"
            try:
                text = blob.decode("utf-8", errors="replace")
            except Exception:
                text = ""
        elif _looks_like_image(mime, name):
            return SummarizeResponse(
                name=name,
                mime_type=mime,
                size_bytes=len(blob),
                extracted_chars=0,
                extractor="fallback",
                truncated=False,
                summary=(
                    f"**{name}** — image, "
                    f"{round(len(blob) / 1024)} KB.\n\n"
                    "Image OCR isn't wired up yet. To summarize what's "
                    "in the image, paste the relevant text directly or "
                    "share the document the image came from."
                ),
            )
        else:
            return SummarizeResponse(
                name=name,
                mime_type=mime,
                size_bytes=len(blob),
                extracted_chars=0,
                extractor="fallback",
                truncated=False,
                summary=(
                    f"**{name}** — {mime or 'binary file'}, "
                    f"{round(len(blob) / 1024)} KB.\n\n"
                    "I don't have an extractor for this format yet. "
                    "Send a `.pdf`, `.docx`, `.md`, `.txt`, `.csv`, "
                    "`.json`, or a code file and I'll summarize the "
                    "actual contents."
                ),
            )
    except Exception:
        log.exception("file extraction failed for %s (%s)", name, mime)
        raise HTTPException(status_code=500, detail="extraction_failed")

    if not text.strip():
        return SummarizeResponse(
            name=name,
            mime_type=mime,
            size_bytes=len(blob),
            extracted_chars=0,
            extractor=extractor,
            truncated=False,
            summary=(
                f"**{name}** — extracted 0 characters of text. The file "
                "might be image-based (scanned PDF), encrypted, or use "
                "an unusual encoding. If it's a scanned PDF, OCR is "
                "needed first."
            ),
        )

    truncated = len(text.encode("utf-8")) > MAX_TEXT_BYTES
    if truncated:
        # Trim from the end (keep the start where summaries / TOCs usually live).
        text = text.encode("utf-8")[:MAX_TEXT_BYTES].decode("utf-8", errors="ignore")

    # Real LLM call — never fall back to canned text. If the LLM is
    # genuinely unavailable we return an error state, not a fake summary.
    if get_active_provider() == "mock":
        return SummarizeResponse(
            name=name,
            mime_type=mime,
            size_bytes=len(blob),
            extracted_chars=len(text),
            extractor=extractor,
            truncated=truncated,
            summary=(
                f"**{name}** — extracted {len(text):,} characters "
                "successfully but no LLM provider is configured to "
                "summarize. Set ANTHROPIC_API_KEY / GROQ_API_KEY / "
                "CEREBRAS_API_KEY in the environment."
            ),
        )

    prompt = _build_summary_prompt(name, mime, text)
    summary = await call_llm(
        system=prompt,
        user_message="Summarize this file.",
        model_tier="reasoning",
        max_tokens=600,
    )
    if not summary:
        raise HTTPException(status_code=503, detail="llm_unavailable")

    return SummarizeResponse(
        name=name,
        mime_type=mime,
        size_bytes=len(blob),
        extracted_chars=len(text),
        extractor=extractor,
        truncated=truncated,
        summary=summary,
    )
